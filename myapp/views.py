from django.shortcuts import redirect, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseRedirect,HttpResponseNotFound
from .models import Document
from django.contrib.auth.models import User
from .forms import DocumentForm
import pickle
from docx.enum.text import WD_COLOR_INDEX
from catboost import CatBoostClassifier, Pool
import docx
import os


best_fac = ['3_9', '3_5', '4_3']
models = pickle.load(open('myapp/static/models.pickle', 'rb'))
TRESHOLD = 0.75
project_path=os.getcwd()

@login_required
def my_view(request):
    message = 'Upload as many files as you want!'
    # Handle file upload
    if request.method == 'POST':
        current_user=request.user
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            print(request.FILES['docfile'])
            newdoc = Document(docfile=request.FILES['docfile'], user=current_user)
            newdoc.save()
            print(newdoc.docfile.url)
            if 'docx' in newdoc.docfile.url:
                doc = docx.Document(project_path+'/'+newdoc.docfile.url)
                cnt=0
                for paragraph in doc.paragraphs:#TODO: Simultaneous prediction for each paragraph in parallel, to improve processing time
                    indicator = False
                    for fac in best_fac:  
                        par = paragraph.text
                        vectorizer, clf = models[fac]
                        if clf.predict_proba(vectorizer.transform([par]))[0][1]>TRESHOLD:
                            cnt+=1
                            indicator = True
                            break
                    if indicator:
                        for run in paragraph.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW#Highlight the paragraph with the found corruption factor
                allCntOfCorruption=cnt
                doc.save(project_path+'/'+newdoc.docfile.url)
            # Redirect to the document list after POST
            return redirect('my-view')
        else:
            message = 'The form is not valid. Fix the following error:'

    else:
        form = DocumentForm()  # An empty, unbound form

    current_user = request.user
    # Load documents for the list psage
    documents = Document.objects.filter(user__id=current_user.id)
    # Render list page with the documents and the form
    context = {'documents': documents, 'form': form, 'message': message}
    return render(request, 'list.html', context)

def delete(request, id):
    try:
        current_user = request.user
        doc = Document.objects.get(id=id,user__id=current_user.id)
        os.remove(project_path+'/media/'+doc.docfile.name)
        doc.delete()
        return HttpResponseRedirect("/")
    except Document.DoesNotExist:
        return HttpResponseNotFound("<h2>Document not found</h2>")